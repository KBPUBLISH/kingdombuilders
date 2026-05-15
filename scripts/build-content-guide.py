"""Build a printable / editable content catalog for Kingdom Builders Publishing.

Pulls live data from:
  - Godly Kids backend  : https://backendgk2-0.onrender.com
      /api/books         (digital library)
      /api/playlists     (audio series)
      /api/amazon-books  (print catalog on Amazon)

  - Faith Defense backend: https://faithdefence.onrender.com
      /categories
      /collections       (seasons)
      /content           (apologetics dialogues)

Outputs (written to ./guide/):
  - Kingdom-Builders-Content-Guide.html
      A TV-guide style document.  Print to PDF from your browser
      (Cmd+P -> Save as PDF) for a polished printable version.
  - Kingdom-Builders-Content-Guide.docx
      Fully editable in Word / Pages / Google Docs, with cover
      images embedded inline.
  - covers/*.jpg
      Cover thumbnails downloaded and resized once for re-use.

Re-run any time you want to refresh the catalog:

    python3 scripts/build-content-guide.py
"""

from __future__ import annotations

import html
import io
import json
import re
import sys
import textwrap
from pathlib import Path
from typing import Any, Iterable
from urllib.parse import urlparse

import requests
from PIL import Image
from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Inches, Pt, RGBColor

ROOT = Path(__file__).resolve().parent.parent
OUT_DIR = ROOT / "guide"
COVERS_DIR = OUT_DIR / "covers"
HTML_PATH = OUT_DIR / "Kingdom-Builders-Content-Guide.html"
DOCX_PATH = OUT_DIR / "Kingdom-Builders-Content-Guide.docx"

GK_BASE = "https://backendgk2-0.onrender.com"
FD_BASE = "https://faithdefence.onrender.com"

REQUEST_TIMEOUT = 60
COVER_MAX_WIDTH = 480  # px, before downscale for both html + docx

# Tune for print economy: short descriptions = fewer pages.
DESCRIPTION_MAX_CHARS = 140

OUT_DIR.mkdir(parents=True, exist_ok=True)
COVERS_DIR.mkdir(parents=True, exist_ok=True)


# ---------------------------------------------------------------------------
# Fetch helpers
# ---------------------------------------------------------------------------

def get_json(url: str, params: dict[str, Any] | None = None) -> Any:
    print(f"  GET {url}", file=sys.stderr)
    r = requests.get(url, params=params, timeout=REQUEST_TIMEOUT)
    r.raise_for_status()
    return r.json()


def unwrap(data: Any) -> list[dict]:
    if isinstance(data, list):
        return data
    if isinstance(data, dict):
        return data.get("data") or data.get("items") or []
    return []


def resolve_gk_media(url: str | None) -> str:
    """Some Godly Kids assets are stored as relative `/uploads/...` paths."""
    if not url:
        return ""
    if url.startswith(("http://", "https://", "blob:")):
        return url
    if url.startswith("/uploads"):
        return f"{GK_BASE}{url}"
    return f"{GK_BASE}/{url.lstrip('/')}"


def book_cover(book: dict) -> str:
    files = book.get("files") or {}
    return resolve_gk_media(book.get("coverImage") or files.get("coverImage"))


def amazon_cover(book: dict) -> str:
    return resolve_gk_media(book.get("coverImage") or (book.get("images") or [""])[0])


def playlist_cover(p: dict) -> str:
    return resolve_gk_media(p.get("coverImage"))


# ---------------------------------------------------------------------------
# Image handling
# ---------------------------------------------------------------------------

def safe_id(text: str) -> str:
    return re.sub(r"[^a-zA-Z0-9_-]", "_", text)[:80] or "cover"


def download_cover(url: str, basename: str) -> Path | None:
    """Download + resize a cover to COVERS_DIR/{basename}.jpg. Returns path or None."""
    if not url:
        return None
    target = COVERS_DIR / f"{basename}.jpg"
    if target.exists() and target.stat().st_size > 0:
        return target
    try:
        r = requests.get(url, timeout=REQUEST_TIMEOUT)
        r.raise_for_status()
        img = Image.open(io.BytesIO(r.content))
        if img.mode in ("RGBA", "P", "LA"):
            bg = Image.new("RGB", img.size, "white")
            bg.paste(img, mask=img.split()[-1] if "A" in img.getbands() else None)
            img = bg
        else:
            img = img.convert("RGB")
        w, h = img.size
        if w > COVER_MAX_WIDTH:
            ratio = COVER_MAX_WIDTH / w
            img = img.resize((COVER_MAX_WIDTH, int(h * ratio)), Image.LANCZOS)
        img.save(target, "JPEG", quality=82, optimize=True)
        return target
    except Exception as e:
        print(f"    ! cover failed ({urlparse(url).netloc}): {e}", file=sys.stderr)
        return None


# ---------------------------------------------------------------------------
# Data fetching
# ---------------------------------------------------------------------------

def fetch_everything() -> dict[str, list[dict]]:
    print("Fetching Godly Kids backend…", file=sys.stderr)
    books = unwrap(get_json(f"{GK_BASE}/api/books", {"status": "published", "limit": 200}))
    playlists = unwrap(get_json(f"{GK_BASE}/api/playlists", {"limit": 200}))
    amazon = unwrap(get_json(f"{GK_BASE}/api/amazon-books", {"status": "published", "limit": 200}))

    print("Fetching Faith Defense backend…", file=sys.stderr)
    fd_categories = unwrap(get_json(f"{FD_BASE}/categories"))
    fd_collections = unwrap(get_json(f"{FD_BASE}/collections"))
    fd_content = unwrap(get_json(f"{FD_BASE}/content", {"limit": 200}))

    return {
        "gk_books": books,
        "gk_playlists": playlists,
        "gk_amazon": amazon,
        "fd_categories": fd_categories,
        "fd_collections": fd_collections,
        "fd_content": fd_content,
    }


# ---------------------------------------------------------------------------
# Normalisation -> Item dataclass (rendered uniformly in both HTML + DOCX)
# ---------------------------------------------------------------------------

def shorten(text: str | None, max_chars: int = DESCRIPTION_MAX_CHARS) -> str:
    """Trim a description down to ~max_chars at a word boundary."""
    if not text:
        return ""
    cleaned = re.sub(r"\s+", " ", text).strip()
    if len(cleaned) <= max_chars:
        return cleaned
    cut = cleaned[: max_chars + 1]
    # Prefer breaking at the last sentence end inside the window.
    for end in [". ", "! ", "? "]:
        idx = cut.rfind(end)
        if idx >= 60:
            return cleaned[: idx + 1]
    # Fall back to a word boundary.
    space = cut.rfind(" ")
    if space > 0:
        return cleaned[:space].rstrip(",.;:") + "…"
    return cleaned[:max_chars].rstrip(",.;:") + "…"


def fmt_duration(seconds: float | int | None) -> str:
    if not seconds:
        return ""
    total = int(seconds)
    h, rem = divmod(total, 3600)
    m, s = divmod(rem, 60)
    if h:
        return f"{h}h {m}m"
    if m:
        return f"{m} min"
    return f"{s} sec"


def build_items(payload: dict[str, list[dict]]) -> dict[str, list[dict]]:
    """Convert raw API rows into a normalised display item per section."""

    # Godly Kids · digital library --------------------------------------------------
    gk_books = []
    for b in payload["gk_books"]:
        meta = []
        cats = b.get("categories") or ([b["category"]] if b.get("category") else [])
        if cats:
            meta.append(("Category", ", ".join(cats[:2])))
        if b.get("minAge"):
            meta.append(("Age", f"{b['minAge']}+"))
        if b.get("readCount"):
            meta.append(("Reads", f"{b['readCount']:,}"))
        if b.get("bookType"):
            meta.append(("Type", str(b["bookType"]).title()))
        gk_books.append(
            {
                "id": str(b.get("_id")),
                "title": b.get("title") or "Untitled",
                "subtitle": (f"by {b['author']}" if b.get("author") else ""),
                "description": shorten(b.get("description")),
                "cover": book_cover(b),
                "meta": meta,
                "badges": [t for t in (
                    "Featured" if b.get("isFeatured") else None,
                    "Members" if b.get("isMembersOnly") else None,
                ) if t],
            }
        )

    # Godly Kids · audio series -----------------------------------------------------
    gk_audio = []
    for p in payload["gk_playlists"]:
        items = p.get("items") or []
        ep = len(items)
        plays = sum((i.get("playCount") or 0) for i in items)
        meta = []
        if ep:
            meta.append(("Episodes", str(ep)))
        cats = p.get("categories") or ([p["category"]] if p.get("category") else [])
        if cats:
            meta.append(("Category", ", ".join(cats[:2])))
        if p.get("type"):
            meta.append(("Type", str(p["type"]).title()))
        if plays:
            meta.append(("Plays", f"{plays:,}"))
        gk_audio.append(
            {
                "id": str(p.get("_id")),
                "title": p.get("title") or "Untitled",
                "subtitle": (f"by {p['author']}" if p.get("author") else ""),
                "description": shorten(p.get("description")),
                "cover": playlist_cover(p),
                "meta": meta,
                "badges": ["Featured"] if p.get("isFeatured") else [],
            }
        )

    # Godly Kids · Amazon (print) ---------------------------------------------------
    gk_amazon = []
    for a in payload["gk_amazon"]:
        meta = []
        if a.get("price"):
            meta.append(("Price", str(a["price"])))
        cats = a.get("categories") or ([a["category"]] if a.get("category") else [])
        if cats:
            meta.append(("Category", ", ".join(cats[:2])))
        if a.get("badgeText"):
            meta.append(("Badge", a["badgeText"]))
        gk_amazon.append(
            {
                "id": str(a.get("_id")),
                "title": a.get("title") or "Untitled",
                "subtitle": (f"by {a['author']}" if a.get("author") else ""),
                "description": shorten(a.get("description")),
                "cover": amazon_cover(a),
                "meta": meta,
                "badges": ["Featured"] if a.get("isFeatured") else [],
                "link": a.get("amazonUrl"),
            }
        )

    # Faith Defense -----------------------------------------------------------------
    fd_categories = [
        {
            "id": c["_id"],
            "title": c.get("name") or "Untitled",
            "subtitle": "",
            "description": "",
            "cover": c.get("coverImageUrl") or "",
            "meta": [("Icon", c.get("icon") or "")],
            "badges": [],
            "_color": c.get("color"),
        }
        for c in payload["fd_categories"]
    ]

    fd_collections = [
        {
            "id": c["_id"],
            "title": c.get("name") or "",
            "subtitle": c.get("subtitle") or "",
            "description": "",
            "cover": "",
            "meta": [],
            "badges": [],
        }
        for c in payload["fd_collections"]
    ]

    fd_dialogues = []
    for c in payload["fd_content"]:
        cat = c.get("categoryId") or {}
        if isinstance(cat, str):
            cat = {}
        meta = []
        if cat.get("name"):
            meta.append(("Topic", cat["name"]))
        if c.get("author"):
            meta.append(("Author", c["author"]))
        if c.get("duration"):
            meta.append(("Length", fmt_duration(c["duration"])))
        if c.get("narrationUrl"):
            meta.append(("Format", "Narrated audio dialogue"))
        else:
            meta.append(("Format", "Dialogue script"))
        fd_dialogues.append(
            {
                "id": str(c.get("_id")),
                "title": c.get("title") or "Untitled",
                "subtitle": cat.get("name") or "",
                "description": shorten(c.get("description")),
                "cover": c.get("coverImageUrl") or "",
                "meta": meta,
                "badges": [cat.get("name")] if cat.get("name") else [],
            }
        )

    return {
        "gk_books": gk_books,
        "gk_audio": gk_audio,
        "gk_amazon": gk_amazon,
        "fd_categories": fd_categories,
        "fd_collections": fd_collections,
        "fd_dialogues": fd_dialogues,
    }


def hydrate_covers(items: Iterable[dict], prefix: str) -> None:
    for it in items:
        url = it.get("cover") or ""
        if not url:
            it["cover_path"] = None
            continue
        basename = f"{prefix}_{safe_id(it['id'])}"
        it["cover_path"] = download_cover(url, basename)


# ---------------------------------------------------------------------------
# HTML rendering
# ---------------------------------------------------------------------------

HTML_CSS = """
:root {
  --ink: #16122a;
  --ink-soft: #524689;
  --ink-mute: #6b6385;
  --gold: #c98f27;
  --gold-soft: #f5edcf;
  --parchment: #fbf7ef;
  --line: rgba(36, 31, 61, 0.12);
}

* { box-sizing: border-box; }

html, body {
  margin: 0;
  padding: 0;
  background: var(--parchment);
  color: var(--ink);
  font-family: "Inter", -apple-system, BlinkMacSystemFont, "Segoe UI", sans-serif;
  font-size: 12.5pt;
  line-height: 1.5;
}

.page {
  max-width: 8.5in;
  margin: 0 auto;
  padding: 0.7in 0.7in 1in;
}

h1, h2, h3, h4 {
  font-family: "Fraunces", Georgia, "Times New Roman", serif;
  color: var(--ink);
  letter-spacing: -0.01em;
  margin: 0;
}

h1 { font-size: 38pt; line-height: 1.05; }
h2 { font-size: 24pt; line-height: 1.1; }
h3 { font-size: 14pt; line-height: 1.2; }

.cover-page {
  min-height: 9.5in;
  padding: 1.5in 0.7in 0.7in;
  background:
    radial-gradient(120% 60% at 50% 0%, rgba(218,171,63,0.22), transparent 60%),
    linear-gradient(180deg, #fbf7ef, #f3eddd);
  page-break-after: always;
}

.cover-page .eyebrow {
  text-transform: uppercase;
  letter-spacing: 0.22em;
  font-size: 9pt;
  color: var(--gold);
  font-weight: 700;
}

.cover-page h1 {
  margin-top: 0.5in;
  font-size: 52pt;
}

.cover-page .subtitle {
  margin-top: 0.25in;
  font-size: 16pt;
  color: var(--ink-soft);
  max-width: 6in;
}

.cover-page .summary {
  margin-top: 1in;
  display: grid;
  grid-template-columns: 1fr 1fr;
  gap: 0.4in;
}

.summary .stat {
  border-left: 3px solid var(--gold);
  padding-left: 0.3in;
}
.summary .stat .num {
  font-family: "Fraunces", serif;
  font-size: 32pt;
  line-height: 1;
}
.summary .stat .lbl {
  text-transform: uppercase;
  letter-spacing: 0.18em;
  font-size: 9pt;
  color: var(--ink-mute);
  margin-top: 6pt;
}

.section {
  padding-top: 0.6in;
  page-break-before: always;
}

.section-header {
  border-bottom: 1px solid var(--line);
  padding-bottom: 0.25in;
  margin-bottom: 0.35in;
}
.section-header .eyebrow {
  text-transform: uppercase;
  letter-spacing: 0.22em;
  font-size: 9pt;
  color: var(--gold);
  font-weight: 700;
}
.section-header h2 {
  margin-top: 6pt;
}
.section-header .lede {
  margin-top: 10pt;
  font-size: 12pt;
  color: var(--ink-mute);
  max-width: 6in;
}

.subhead {
  text-transform: uppercase;
  letter-spacing: 0.2em;
  font-size: 10pt;
  font-weight: 700;
  color: var(--ink-soft);
  margin: 0.35in 0 0.15in;
  display: flex;
  justify-content: space-between;
  align-items: baseline;
}
.subhead .count {
  font-size: 9pt;
  color: var(--ink-mute);
  font-weight: 600;
}

.chips {
  display: flex;
  flex-wrap: wrap;
  gap: 6pt;
  margin: 0 0 0.25in;
}
.chip {
  display: inline-flex;
  align-items: center;
  gap: 4pt;
  background: var(--ink);
  color: var(--parchment);
  padding: 4pt 10pt;
  border-radius: 999px;
  font-size: 9pt;
}
.chip .accent { color: var(--gold-soft); font-weight: 600; }

.grid-2 {
  display: grid;
  grid-template-columns: 1fr 1fr;
  gap: 0.18in 0.22in;
}

.item {
  display: grid;
  grid-template-columns: 1.05in 1fr;
  gap: 0.14in;
  padding: 0.1in 0;
  border-bottom: 1px solid var(--line);
  page-break-inside: avoid;
  align-items: start;
}
.item .cover {
  width: 1.0in;
  border-radius: 6pt;
  overflow: hidden;
  border: 1px solid var(--line);
  align-self: start;
  line-height: 0;
}
.item .cover img {
  width: 100%;
  display: block;
}
.item .cover.placeholder {
  line-height: 1.2;
  height: 1.35in;
  display: flex;
  align-items: center;
  justify-content: center;
  font-family: "Fraunces", serif;
  font-size: 9pt;
  color: var(--ink-mute);
  text-align: center;
  padding: 6pt;
  background: #ece8db;
}

.item .body h3 { font-size: 11.5pt; line-height: 1.15; }
.item .body .sub {
  font-style: italic;
  color: var(--ink-soft);
  margin-top: 2pt;
  font-size: 8.5pt;
}
.item .body p.desc {
  margin: 5pt 0 4pt;
  color: var(--ink);
  font-size: 9pt;
  line-height: 1.35;
}
.item .body .badges {
  display: flex; flex-wrap: wrap; gap: 3pt;
  margin-top: 4pt;
}
.item .body .badge {
  background: var(--gold-soft);
  color: var(--ink);
  font-size: 7pt;
  text-transform: uppercase;
  letter-spacing: 0.1em;
  padding: 1pt 6pt;
  border-radius: 999px;
  font-weight: 700;
}
.item .body .meta {
  margin-top: 5pt;
  display: flex;
  flex-wrap: wrap;
  gap: 2pt 12pt;
  font-size: 8pt;
  color: var(--ink-mute);
}
.item .body .meta .row {
  display: inline-flex;
  gap: 4pt;
}
.item .body .meta .k {
  font-weight: 700;
  color: var(--ink-soft);
  text-transform: uppercase;
  letter-spacing: 0.12em;
  font-size: 7pt;
}

.footer {
  border-top: 1px solid var(--line);
  margin-top: 0.5in;
  padding-top: 0.2in;
  font-size: 9pt;
  color: var(--ink-mute);
}

@media print {
  body { background: white; }
  .page { padding: 0; max-width: none; }
  .item .body details.episodes[open] summary { display: none; }
  .item .body details.episodes ol { display: block; }
}
"""


def html_meta_grid(meta: list[tuple[str, str]]) -> str:
    if not meta:
        return ""
    rows = "".join(
        f'<span class="row"><span class="k">{html.escape(k)}</span>'
        f'<span class="v">{html.escape(v)}</span></span>'
        for k, v in meta
        if v
    )
    return f'<div class="meta">{rows}</div>'


def html_item(item: dict) -> str:
    cover_path = item.get("cover_path")
    if cover_path:
        cover_html = (
            f'<div class="cover"><img src="covers/{cover_path.name}" alt=""></div>'
        )
    else:
        cover_html = (
            f'<div class="cover placeholder">{html.escape(item.get("title",""))}</div>'
        )
    badges_html = ""
    if item.get("badges"):
        badges_html = '<div class="badges">' + "".join(
            f'<span class="badge">{html.escape(b)}</span>' for b in item["badges"]
        ) + "</div>"

    desc = item.get("description") or ""
    desc_html = f'<p class="desc">{html.escape(desc)}</p>' if desc else ""

    sub = item.get("subtitle") or ""
    sub_html = f'<p class="sub">{html.escape(sub)}</p>' if sub else ""

    return f"""
    <article class="item">
      {cover_html}
      <div class="body">
        <h3>{html.escape(item['title'])}</h3>
        {sub_html}
        {badges_html}
        {desc_html}
        {html_meta_grid(item['meta'])}
      </div>
    </article>
    """.strip()


def render_html(buckets: dict[str, list[dict]]) -> str:
    total_titles = (
        len(buckets["gk_books"])
        + len(buckets["gk_audio"])
        + len(buckets["gk_amazon"])
        + len(buckets["fd_dialogues"])
    )

    chips = "".join(
        f'<span class="chip"><span class="accent">{html.escape(c["title"])}</span>'
        + (f" · {html.escape(c['subtitle'])}" if c.get("subtitle") else "")
        + "</span>"
        for c in buckets["fd_collections"]
    )

    fd_categories_chips = "".join(
        f'<span class="chip" style="background:{html.escape(c.get("_color") or "#16122a")}">'
        f'<span class="accent">{html.escape(c.get("meta", [["",""]])[0][1] or "")}</span>'
        f' {html.escape(c["title"])}</span>'
        for c in buckets["fd_categories"]
    )

    def section(items: list[dict], title: str) -> str:
        if not items:
            return ""
        body = "\n".join(html_item(it) for it in items)
        return (
            f'<div class="subhead"><span>{html.escape(title)}</span>'
            f'<span class="count">{len(items)} items</span></div>\n'
            f'<div class="grid-2">{body}</div>'
        )

    return f"""<!DOCTYPE html>
<html lang="en">
<head>
<meta charset="utf-8">
<title>Kingdom Builders Publishing · Content Catalog</title>
<link rel="preconnect" href="https://fonts.googleapis.com">
<link rel="preconnect" href="https://fonts.gstatic.com" crossorigin>
<link href="https://fonts.googleapis.com/css2?family=Fraunces:opsz,wght@9..144,400;9..144,600;9..144,700&family=Inter:wght@400;500;600;700&display=swap" rel="stylesheet">
<style>{HTML_CSS}</style>
</head>
<body>
<div class="page cover-page">
  <p class="eyebrow">Kingdom Builders Publishing</p>
  <h1>Content Catalog</h1>
  <p class="subtitle">Every published book, audio series, and apologetics dialogue from the Godly Kids and Faith Defense apps — pulled live from our backends.</p>

  <div class="summary">
    <div class="stat"><div class="num">{len(buckets['gk_books'])}</div><div class="lbl">Digital Books</div></div>
    <div class="stat"><div class="num">{len(buckets['gk_audio'])}</div><div class="lbl">Audio Series</div></div>
    <div class="stat"><div class="num">{len(buckets['gk_amazon'])}</div><div class="lbl">Print Books on Amazon</div></div>
    <div class="stat"><div class="num">{len(buckets['fd_dialogues'])}</div><div class="lbl">Faith Defense Dialogues</div></div>
  </div>

  <div class="footer">{total_titles} titles · Generated automatically — re-run <code>python3 scripts/build-content-guide.py</code> to refresh.</div>
</div>

<div class="page section">
  <header class="section-header">
    <p class="eyebrow">App 01 · Godly Kids</p>
    <h2>The Godly Kids library</h2>
    <p class="lede">A Christian library for kids and families — devotionals, original stories, audio adventures, and a growing print catalog on Amazon.</p>
  </header>

  {section(buckets['gk_books'], 'Digital books')}
  {section(buckets['gk_audio'], 'Audio series')}
  {section(buckets['gk_amazon'], 'Print catalog · Amazon')}
</div>

<div class="page section">
  <header class="section-header">
    <p class="eyebrow">App 02 · Faith Defense</p>
    <h2>Faith Defense — apologetics for the next generation</h2>
    <p class="lede">Original dialogues that teach teens and curious minds how to defend their faith with honest answers about science, history, culture, and the Bible.</p>
  </header>

  <div class="subhead"><span>Topics ({len(buckets['fd_categories'])})</span></div>
  <div class="chips">{fd_categories_chips}</div>

  <div class="subhead"><span>Seasons ({len(buckets['fd_collections'])})</span></div>
  <div class="chips">{chips}</div>

  {section(buckets['fd_dialogues'], 'Dialogues')}

  <div class="footer">Kingdom Builders Publishing · kbpublish.org</div>
</div>

</body>
</html>
""".strip()


# ---------------------------------------------------------------------------
# DOCX rendering
# ---------------------------------------------------------------------------

INK = RGBColor(0x16, 0x12, 0x2A)
INK_SOFT = RGBColor(0x52, 0x46, 0x89)
INK_MUTE = RGBColor(0x6B, 0x63, 0x85)
GOLD = RGBColor(0xC9, 0x8F, 0x27)


def set_cell_shading(cell, color_hex: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), color_hex)
    tc_pr.append(shd)


def add_eyebrow(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    run = p.add_run(text.upper())
    run.font.name = "Inter"
    run.font.size = Pt(9)
    run.font.color.rgb = GOLD
    run.bold = True
    p.paragraph_format.space_after = Pt(2)


def add_title(doc: Document, text: str, size: int = 28) -> None:
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = "Fraunces"
    run.font.size = Pt(size)
    run.bold = True
    run.font.color.rgb = INK
    p.paragraph_format.space_after = Pt(10)


def add_lede(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = "Inter"
    run.font.size = Pt(11)
    run.font.color.rgb = INK_MUTE
    p.paragraph_format.space_after = Pt(18)


def add_subhead(doc: Document, text: str, count: int | None = None) -> None:
    p = doc.add_paragraph()
    run = p.add_run(text.upper())
    run.font.name = "Inter"
    run.font.size = Pt(10)
    run.bold = True
    run.font.color.rgb = INK_SOFT
    if count is not None:
        tail = p.add_run(f"   ({count})")
        tail.font.name = "Inter"
        tail.font.size = Pt(9)
        tail.font.color.rgb = INK_MUTE
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after = Pt(8)


def _clear_cell(cell) -> None:
    for p in list(cell.paragraphs):
        p._element.getparent().remove(p._element)


def add_item_into(cell, item: dict) -> None:
    """Render a single item compactly inside a cell.

    Layout: a small inner table — cover (left, ~0.85in) + body (right) —
    so cards stack nicely two-per-row in the outer grid.
    """
    _clear_cell(cell)
    inner = cell.add_table(rows=1, cols=2)
    inner.autofit = False
    inner.columns[0].width = Inches(0.85)
    inner.columns[1].width = Inches(2.4)
    cover_cell = inner.rows[0].cells[0]
    body_cell = inner.rows[0].cells[1]
    cover_cell.width = Inches(0.85)
    body_cell.width = Inches(2.4)
    cover_cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP
    body_cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP

    cover_path = item.get("cover_path")
    cp = cover_cell.paragraphs[0]
    cp.paragraph_format.space_after = Pt(0)
    if cover_path and Path(cover_path).exists():
        cp.add_run().add_picture(str(cover_path), width=Inches(0.8))
    else:
        run = cp.add_run(item["title"])
        run.font.name = "Fraunces"
        run.font.size = Pt(8)
        run.italic = True
        run.font.color.rgb = INK_MUTE
        set_cell_shading(cover_cell, "F1ECDC")

    _clear_cell(body_cell)

    title_p = body_cell.add_paragraph()
    title_p.paragraph_format.space_after = Pt(1)
    tr = title_p.add_run(item["title"])
    tr.font.name = "Fraunces"
    tr.font.size = Pt(10.5)
    tr.bold = True
    tr.font.color.rgb = INK

    if item.get("subtitle"):
        sp = body_cell.add_paragraph()
        sp.paragraph_format.space_after = Pt(2)
        sr = sp.add_run(item["subtitle"])
        sr.italic = True
        sr.font.name = "Inter"
        sr.font.size = Pt(8)
        sr.font.color.rgb = INK_SOFT

    if item.get("badges"):
        bp = body_cell.add_paragraph()
        bp.paragraph_format.space_after = Pt(2)
        for i, b in enumerate(item["badges"]):
            if i:
                bp.add_run(" · ").font.color.rgb = INK_MUTE
            br = bp.add_run(b.upper())
            br.bold = True
            br.font.name = "Inter"
            br.font.size = Pt(7)
            br.font.color.rgb = GOLD

    desc = (item.get("description") or "").strip()
    if desc:
        dp = body_cell.add_paragraph()
        dp.paragraph_format.space_after = Pt(2)
        dr = dp.add_run(desc)
        dr.font.name = "Inter"
        dr.font.size = Pt(8.5)
        dr.font.color.rgb = INK

    if item.get("meta"):
        mp = body_cell.add_paragraph()
        mp.paragraph_format.space_after = Pt(0)
        first = True
        for k, v in item["meta"]:
            if not v:
                continue
            if not first:
                mp.add_run("  ·  ").font.color.rgb = INK_MUTE
            kr = mp.add_run(f"{k}: ")
            kr.bold = True
            kr.font.name = "Inter"
            kr.font.size = Pt(7.5)
            kr.font.color.rgb = INK_SOFT
            vr = mp.add_run(str(v))
            vr.font.name = "Inter"
            vr.font.size = Pt(7.5)
            vr.font.color.rgb = INK_MUTE
            first = False


def add_items_grid(doc: Document, items: list[dict]) -> None:
    """Render items two-per-row inside an outer 2-column table."""
    if not items:
        return
    rows = (len(items) + 1) // 2
    outer = doc.add_table(rows=rows, cols=2)
    outer.autofit = False
    outer.columns[0].width = Inches(3.4)
    outer.columns[1].width = Inches(3.4)
    for idx, it in enumerate(items):
        r, c = divmod(idx, 2)
        cell = outer.rows[r].cells[c]
        cell.width = Inches(3.4)
        add_item_into(cell, it)
    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_after = Pt(4)


def render_docx(buckets: dict[str, list[dict]], path: Path) -> None:
    doc = Document()

    # Page margins
    for section in doc.sections:
        section.top_margin = Inches(0.7)
        section.bottom_margin = Inches(0.7)
        section.left_margin = Inches(0.7)
        section.right_margin = Inches(0.7)

    # Cover --------------------------------------------------------------
    add_eyebrow(doc, "Kingdom Builders Publishing")
    add_title(doc, "Content Catalog", size=36)
    add_lede(
        doc,
        "Every published book, audio series, and apologetics dialogue from the "
        "Godly Kids and Faith Defense apps — pulled live from our backends.",
    )

    summary = doc.add_table(rows=2, cols=2)
    summary.autofit = True
    cells = [
        ("Digital books", len(buckets["gk_books"])),
        ("Audio series", len(buckets["gk_audio"])),
        ("Print on Amazon", len(buckets["gk_amazon"])),
        ("Faith Defense dialogues", len(buckets["fd_dialogues"])),
    ]
    for i, (label, n) in enumerate(cells):
        cell = summary.rows[i // 2].cells[i % 2]
        p = cell.paragraphs[0]
        nr = p.add_run(f"{n}")
        nr.bold = True
        nr.font.size = Pt(22)
        nr.font.name = "Fraunces"
        nr.font.color.rgb = INK
        p.add_run("\n")
        lr = p.add_run(label.upper())
        lr.bold = True
        lr.font.size = Pt(8)
        lr.font.name = "Inter"
        lr.font.color.rgb = INK_MUTE

    doc.add_page_break()

    # Section: Godly Kids -------------------------------------------------
    add_eyebrow(doc, "App 01 · Godly Kids")
    add_title(doc, "The Godly Kids library", size=24)
    add_lede(
        doc,
        "A Christian library for kids and families — devotionals, original stories, "
        "audio adventures, and a growing print catalog on Amazon.",
    )

    if buckets["gk_books"]:
        add_subhead(doc, "Digital books", len(buckets["gk_books"]))
        add_items_grid(doc, buckets["gk_books"])

    if buckets["gk_audio"]:
        add_subhead(doc, "Audio series", len(buckets["gk_audio"]))
        add_items_grid(doc, buckets["gk_audio"])

    if buckets["gk_amazon"]:
        add_subhead(doc, "Print catalog · Amazon", len(buckets["gk_amazon"]))
        add_items_grid(doc, buckets["gk_amazon"])

    doc.add_page_break()

    # Section: Faith Defense ---------------------------------------------
    add_eyebrow(doc, "App 02 · Faith Defense")
    add_title(doc, "Faith Defense — apologetics for the next generation", size=22)
    add_lede(
        doc,
        "Original dialogues that teach teens and curious minds how to defend their "
        "faith with honest answers about science, history, culture, and the Bible.",
    )

    if buckets["fd_categories"]:
        add_subhead(doc, "Topics", len(buckets["fd_categories"]))
        add_items_grid(doc, buckets["fd_categories"])

    if buckets["fd_collections"]:
        add_subhead(doc, "Seasons", len(buckets["fd_collections"]))
        add_items_grid(doc, buckets["fd_collections"])

    if buckets["fd_dialogues"]:
        add_subhead(doc, "Dialogues", len(buckets["fd_dialogues"]))
        add_items_grid(doc, buckets["fd_dialogues"])

    # Footer note
    footer = doc.add_paragraph()
    fr = footer.add_run(
        "Generated automatically · rerun python3 scripts/build-content-guide.py to refresh."
    )
    fr.font.size = Pt(8)
    fr.font.color.rgb = INK_MUTE

    doc.save(str(path))


# ---------------------------------------------------------------------------
# Main
# ---------------------------------------------------------------------------

def main() -> int:
    payload = fetch_everything()
    buckets = build_items(payload)

    print("\nDownloading covers…", file=sys.stderr)
    hydrate_covers(buckets["gk_books"], "gkb")
    hydrate_covers(buckets["gk_audio"], "gka")
    hydrate_covers(buckets["gk_amazon"], "gkz")
    hydrate_covers(buckets["fd_categories"], "fdc")
    hydrate_covers(buckets["fd_collections"], "fds")
    hydrate_covers(buckets["fd_dialogues"], "fdd")

    print("\nWriting HTML…", file=sys.stderr)
    HTML_PATH.write_text(render_html(buckets), encoding="utf-8")
    print(f"  -> {HTML_PATH.relative_to(ROOT)}", file=sys.stderr)

    print("\nWriting DOCX…", file=sys.stderr)
    render_docx(buckets, DOCX_PATH)
    print(f"  -> {DOCX_PATH.relative_to(ROOT)}", file=sys.stderr)

    total = sum(
        len(v) for k, v in buckets.items() if k not in ("fd_categories", "fd_collections")
    )
    print(
        f"\nDone. {total} titles catalogued across both apps.",
        file=sys.stderr,
    )
    return 0


if __name__ == "__main__":
    sys.exit(main())
